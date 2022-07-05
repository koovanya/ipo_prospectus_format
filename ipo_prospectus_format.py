#依据《中国证券监督管理委员会公告[2020]39号——公开发行证券的公司信息披露内容与格式准则第45号——科创板上市公司发行证券申请文件》编写
#原文件地址http://www.sse.com.cn/lawandrules/regulations/csrcannoun/c/5178163.pdf

from docx import Document
import docx.shared
import openpyxl
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


#创建新word
doc = Document()

#导入目录模板（excel文件）
workbook = openpyxl.load_workbook('template.xlsx')
sheet = workbook['content']

#将目录模板（excel文件）的数据取出，放到列表里
rows = sheet.rows
lst = []
for row in rows:
    row_lst = []
    for cell in row:
        row_lst.append(cell.value)
    lst.append(row_lst)

#遍历每一行
for i in range(0,sheet.max_row):
        #若第一列不为空，则设置该列文字为一级标题
        if lst[i][0] != None:
            title = doc.add_heading(lst[i][0],level=1)
            #给一级标题设置样式
            for run in title.runs:
                #设置字号（三号）
                run.font.size = docx.shared.Pt(16)
                #取消加粗
                run.font.bold = False
                #设置字体（黑体）
                run.font.name = '黑体'
                #设置颜色（黑）
                run.font.color.rgb=docx.shared.RGBColor(0,0,0)
                #设置东亚字体
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), '黑体')
                #设置一级标题居中
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # 设置1.5倍行距
                title.paragraph_format.line_spacing = 1.5
        #若第一列为空，则设置第二列文字为二级标题
        if lst[i][0] is None:
            title = doc.add_heading(lst[i][1],level=2)
            #给二级标题设置样式
            for run in title.runs:
                #设置字号（四号）
                run.font.size = docx.shared.Pt(14)
                #取消加粗
                run.font.bold = False
                #设置字体（黑体）
                run.font.name = '黑体'
                #设置颜色（黑）
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                #设置东亚字体
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), '黑体')
                # 设置1.5倍行距
                title.paragraph_format.line_spacing = 1.5


                #插入正文样本
                para = doc.add_paragraph()
                run2 = para.add_run('正文样本\n')
                #设置正文字号（小四）
                run2.font.size = docx.shared.Pt(12)
                #取消加粗
                run2.font.bold = False
                #设置字体（宋体）
                run2.font.name = '宋体'
                #设置颜色（黑）
                run2.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                #设置东亚字体
                r = run2._element.rPr.rFonts
                r.set(qn('w:eastAsia'), '宋体')
                # 设置1.5倍行距
                para.paragraph_format.line_spacing = 1.5

#保存文件
doc.save('ipo_prospectus_format.docx')

